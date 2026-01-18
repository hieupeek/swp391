
import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re

def create_docx_from_md(md_path, docx_path):
    doc = docx.Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    current_table = []
    in_table = False

    for line in lines:
        line = line.strip()

        # Handle Tables
        if line.startswith('|'):
            if not in_table:
                in_table = True
                current_table = []
            
            # Remove leading/trailing pipes and split
            row_content = [cell.strip() for cell in line.strip('|').split('|')]
            # Remove separator lines like |---|---|
            if '---' in row_content[0]:
                continue
            current_table.append(row_content)
            continue
        else:
            if in_table:
                # Flush table to doc
                if current_table:
                    # Determine max columns
                    max_cols = max(len(row) for row in current_table)
                    table = doc.add_table(rows=len(current_table), cols=max_cols)
                    table.style = 'Table Grid'
                    
                    for i, row_data in enumerate(current_table):
                        row_cells = table.rows[i].cells
                        for j, cell_text in enumerate(row_data):
                            if j < len(row_cells):
                                # Clean bold markdown ** from cell text
                                cell_text_clean = cell_text.replace('**', '')
                                row_cells[j].text = cell_text_clean
                in_table = False
                doc.add_paragraph() # Add spacing after table

        if not line:
            continue
            
        # Handle Headers
        if line.startswith('# '):
            doc.add_heading(line[2:], level=0)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=1)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=2)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=3)
            
        # Handle Lists (Bullet points)
        elif line.startswith('* ') or line.startswith('- '):
            # Check for bold text inside list item
            text = line[2:]
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, text)
            
        # Handle Blockquotes
        elif line.startswith('> '):
            p = doc.add_paragraph(line[2:])
            p.italic = True
            
        # Handle Metadata lines (YAML frontmatter)
        elif line.startswith('---') or ':' in line and not line.startswith('*'):
             # Simplistic skip for YAML frontmatter if it looks like key: value pairs at start
             # But for body text, just adding it
             if not line.startswith('---'):
                 p = doc.add_paragraph(line)
        
        # Normal Paragraph
        else:
             p = doc.add_paragraph()
             add_formatted_text(p, line)

    doc.save(docx_path)
    print(f"Successfully created {docx_path}")

def add_formatted_text(paragraph, text):
    """Parses simple bold markdown (**text**) and adds runs to paragraph."""
    # Split by double asterisks
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)

if __name__ == "__main__":
    md_file = "/Users/vuhieu/Documents/swp391/_bmad-output/analysis/vision-and-scope.md"
    docx_file = "/Users/vuhieu/Documents/swp391/_bmad-output/analysis/vision-and-scope.docx"
    create_docx_from_md(md_file, docx_file)
